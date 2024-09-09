import os
import re
import cv2
import enchant
import pdfkit
import pytesseract
import base64
import requests
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from spellchecker import SpellChecker
from langdetect import detect
from googletrans import Translator
from io import BytesIO
from PIL import Image

matplotlib.use('Agg')

dataPath = r"C:\Users\90553\Desktop\Images\Test"
txtPath = r"C:\Users\90553\Desktop\Images\textextract.txt"

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
custom_config = r'-l eng+tur'


# Eski kod
def pathBul(path):
    list_of_paths = []
    for (dirpath, dirnames, filenames) in os.walk(path):
        for filename in filenames:
            if filename.endswith('.jpg') or filename.endswith('.JPG') or filename.endswith('.png') or filename.endswith(
                    '.PNG') or filename.endswith('.jpeg'):
                list_of_paths.append(os.sep.join([dirpath, filename]))
    return list_of_paths


def ocr(paths):
    for path in paths:
        original_image = cv2.imread(path)
        grayscale_image = cv2.cvtColor(original_image, cv2.COLOR_BGR2GRAY)
        best_readings = []

        for x in range(0, 11):
            increased_size_original_image = changeImageSizeBy5(original_image, 100 + (5 * x))
            decreased_size_original_image = changeImageSizeBy5(original_image, 100 - (5 * x))
            increased_size_grayscale_image = changeImageSizeBy5(grayscale_image, 100 + (5 * x))
            decreased_size_grayscale_image = changeImageSizeBy5(grayscale_image, 100 - (5 * x))

            increased_size_original_image_text = pytesseract.image_to_string(increased_size_original_image,
                                                                             config=custom_config)
            decreased_size_original_image_text = pytesseract.image_to_string(decreased_size_original_image,
                                                                             config=custom_config)
            increased_size_grayscale_image_text = pytesseract.image_to_string(increased_size_grayscale_image,
                                                                              config=custom_config)
            decreased_size_grayscale_image_text = pytesseract.image_to_string(decreased_size_grayscale_image,
                                                                              config=custom_config)

            string_list = [increased_size_original_image_text, decreased_size_original_image_text,
                           increased_size_grayscale_image_text, decreased_size_grayscale_image_text]

            best_readings.append(findLongestString(string_list))

        best_readings = removeUnnecessaryChars(best_readings)
        if len(best_readings) > 0:
            best_reading = max(best_readings, key=len)
        else:
            best_reading = ""

        translated_reading = ""
        language = detectLanguages(best_reading)
        if language == "NOT FOUND":
            print("dil bulunamadi")
        else:
            if language == "tr":
                corrected_reading = turkishTypoCorrection(best_reading)
                translated_reading = translate("en", corrected_reading)
            else:
                corrected_reading = englishTypoCorrection(best_reading)
                translated_reading = translate("tr", corrected_reading)

        with open(txtPath, 'a') as file:
            file.write(best_reading)
            file.write('\n')
            if translated_reading != "" and translated_reading.text is not None and len(translated_reading.text) > 0:
                file.write(translated_reading.text)
                file.write('\n')
            file.write("-------")
            file.write('\n')


# Kötü çalışan kod
def getSingleImage(image_url, lang, is_pdf_or_word):
    if image_url.startswith("data:image/jpeg;base64"):
        url_parts = image_url.split(',')
        base64_image_data = url_parts[1]

        image_data = base64.b64decode(base64_image_data)
        image = Image.open(BytesIO(image_data))

        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    else:
        image_path = image_url.replace("http://127.0.0.1:8000/", "")
        image_cv = cv2.imread(image_path)

    if is_pdf_or_word == "id":
        return ocr_for_identity(image_cv)
    else:
        return ocr2(image_cv, lang, is_pdf_or_word)


def ocr2(original_image, lang, is_pdf_or_word):
    grayscale_image = cv2.cvtColor(original_image, cv2.COLOR_BGR2GRAY)
    best_readings = []

    for x in range(0, 11):
        increased_size_original_image = changeImageSizeBy5(original_image, 100 + (5 * x))
        decreased_size_original_image = changeImageSizeBy5(original_image, 100 - (5 * x))
        increased_size_grayscale_image = changeImageSizeBy5(grayscale_image, 100 + (5 * x))
        decreased_size_grayscale_image = changeImageSizeBy5(grayscale_image, 100 - (5 * x))

        increased_size_original_image_text = pytesseract.image_to_string(increased_size_original_image,
                                                                         config=custom_config)
        decreased_size_original_image_text = pytesseract.image_to_string(decreased_size_original_image,
                                                                         config=custom_config)
        increased_size_grayscale_image_text = pytesseract.image_to_string(increased_size_grayscale_image,
                                                                          config=custom_config)
        decreased_size_grayscale_image_text = pytesseract.image_to_string(decreased_size_grayscale_image,
                                                                          config=custom_config)

        string_list = [increased_size_original_image_text, decreased_size_original_image_text,
                       increased_size_grayscale_image_text, decreased_size_grayscale_image_text]

        best_readings.append(findLongestString(string_list))

    best_readings = removeUnnecessaryChars(best_readings)
    if len(best_readings) > 0:
        best_reading = max(best_readings, key=len)
    else:
        best_reading = ""

    correction_lang = detectLanguages(best_reading)
    if is_pdf_or_word == "text":
        if correction_lang == "NOT FOUND":
            translated_reading = translate(lang, best_reading)
        else:
            if correction_lang == "tr":
                corrected_reading = turkishTypoCorrection(best_reading)
            else:
                corrected_reading = englishTypoCorrection(best_reading)

            translated_reading = translate(lang, corrected_reading)

        if translated_reading != "" and translated_reading.text is not None and len(translated_reading.text) > 0:
            return translated_reading.text
        else:
            return "OKUNAMADI"
    else:
        if correction_lang == "tr":
            corrected_reading = turkishTypoCorrection(best_reading)
            if lang == "none" and is_pdf_or_word == "pdf":
                return create_pdf(corrected_reading)
            elif lang != "none" and is_pdf_or_word == "pdf":
                translated_reading = translate(lang, corrected_reading)
                return create_pdf(translated_reading.text)
            elif lang == "none" and is_pdf_or_word == "word":
                return create_word(corrected_reading)
            elif lang != "none" and is_pdf_or_word == "word":
                translated_reading = translate(lang, corrected_reading)
                return create_word(translated_reading.text)

        elif correction_lang == "en":
            corrected_reading = englishTypoCorrection(best_reading)
            if lang == "none" and is_pdf_or_word == "pdf":
                return create_pdf(corrected_reading)
            elif lang != "none" and is_pdf_or_word == "pdf":
                translated_reading = translate(lang, corrected_reading)
                return create_pdf(translated_reading.text)
            elif lang == "none" and is_pdf_or_word == "word":
                return create_word(corrected_reading)
            elif lang != "none" and is_pdf_or_word == "word":
                translated_reading = translate(lang, corrected_reading)
                return create_word(translated_reading.text)
        else:
            if lang == "none" and is_pdf_or_word == "pdf":
                return create_pdf(best_reading)
            elif lang != "none" and is_pdf_or_word == "pdf":
                translated_reading = translate(lang, best_reading)
                return create_pdf(translated_reading.text)
            elif lang == "none" and is_pdf_or_word == "word":
                return create_word(best_reading)
            elif lang != "none" and is_pdf_or_word == "word":
                translated_reading = translate(lang, best_reading)
                return create_word(translated_reading.text)

    """with open(txtPath, 'a') as file:
        file.write(best_reading)
        file.write('\n')
        if translated_reading != "" and translated_reading.text is not None and len(translated_reading.text) > 0:
            file.write(translated_reading.text)
            file.write('\n')
        file.write("-------")
        file.write('\n')"""


def ocr_for_identity(image):
    increased_size_original_image = changeImageSizeBy5(image, 100 + (5 * 3))
    increased_size_original_image_text = pytesseract.image_to_string(increased_size_original_image)

    return create_word(increased_size_original_image_text)


def changeImageSizeBy5(image, percentage):
    width = int(image.shape[1] * percentage / 100)
    height = int(image.shape[0] * percentage / 100)

    resized_image = cv2.resize(image, (width, height), interpolation=cv2.INTER_LINEAR)
    return resized_image


def findLongestString(string_list):
    if string_list is not None and len(string_list) > 0:
        max_length = 0
        longest_string = ""
        for string in string_list:
            if len(string) > max_length:
                max_length = len(string)
                longest_string = string

        return longest_string

    else:
        return ""


def removeUnnecessaryChars(text):
    if text is not None and len(text) > 0:
        new_readings = []
        temp_readings = []
        final_readings = []
        for line in text:
            modified_text = line.replace("\n", " ")
            cleaned_text = " ".join(modified_text.split())
            new_readings.append(cleaned_text)

        for line in new_readings:
            for char in line:
                if (ord(char) <= 33 or
                        (33 <= ord(char) <= 39) or
                        (42 <= ord(char) <= 45) or
                        (59 <= ord(char) <= 64) or
                        (91 <= ord(char) <= 96) or
                        (123 <= ord(char) <= 127) or
                        (130 <= ord(char) <= 134) or
                        (136 <= ord(char) <= 147) or
                        (149 <= ord(char) <= 152) or
                        (155 <= ord(char) <= 172) or
                        (174 <= ord(char) <= 255) or
                        ord(char) >= 255):
                    final_text = line.replace(char, '')
                    if final_text not in temp_readings:
                        temp_readings.append(final_text)

        for line in temp_readings:
            cleaned_text = " ".join(line.split())
            final_readings.append(cleaned_text)

        return final_readings
    else:
        return ""


def detectLanguages(text):
    if text is not None and len(text) > 0:
        return detect(text)
    else:
        return "NOT FOUND"


def englishTypoCorrection(text):
    if text is not None and len(text) > 0:
        spell = SpellChecker()
        correctedWords = []
        correctedText = ""
        words = text.split(" ")

        for word in words:
            corrected_word = spell.correction(word)
            if corrected_word is not None:
                correctedWords.append(corrected_word)
            else:
                corrected_word_candidate = spell.candidates(word)
                if corrected_word_candidate is not None:
                    correctedWords.append(corrected_word_candidate)
                else:
                    correctedWords.append(word)

        for word in correctedWords:
            correctedText += word + " "

        return correctedText.upper()

    else:
        return ""


def turkishTypoCorrection(text):
    if text is not None and len(text) > 0:
        turkishDict = enchant.Dict("tr_TR")
        correctedWords = []
        correctedText = ""
        words = text.split(" ")

        for word in words:
            if word != '':
                if turkishDict.check(word):
                    correctedWords.append(word)
                else:
                    suggestions = turkishDict.suggest(word)
                    if suggestions:
                        correctedWords.append(suggestions[0])
                    else:
                        correctedWords.append(word)

        for word in correctedWords:
            correctedText += word + " "

        return correctedText.upper()

    else:
        return ""


# Fiş Bulma Ve O kısmı Alma
def getSingleReceiptForFinding(image_url):
    if image_url.startswith("data:image/jpeg;base64"):
        url_parts = image_url.split(',')
        base64_image_data = url_parts[1]

        image_data = base64.b64decode(base64_image_data)
        image = Image.open(BytesIO(image_data))

        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    else:
        image_path = image_url.replace("http://127.0.0.1:8000/", "")
        image_cv = cv2.imread(image_path)

    return findReceipt(image_cv)


def findReceipt(image):
    resize_ratio = 500 / image.shape[0]
    original = image.copy()
    image = opencv_resize(image, resize_ratio)

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    blurred = cv2.GaussianBlur(gray, (5, 5), 0)

    rectKernel = cv2.getStructuringElement(cv2.MORPH_RECT, (9, 9))
    dilated = cv2.dilate(blurred, rectKernel)

    edged = cv2.Canny(dilated, 100, 200, apertureSize=3)

    contours, hierarchy = cv2.findContours(edged, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)

    largest_contours = sorted(contours, key=cv2.contourArea, reverse=True)[:10]

    get_receipt_contour(largest_contours)

    receipt_contour = get_receipt_contour(largest_contours)

    scanned = wrap_perspective(original.copy(), contour_to_rect(receipt_contour, resize_ratio))
    plt.figure(figsize=(16, 10))

    result_file_path = "scanned_receipt.png"
    cv2.imwrite(result_file_path, scanned)
    return result_file_path


def opencv_resize(image, ratio):
    width = int(image.shape[1] * ratio)
    height = int(image.shape[0] * ratio)
    dim = (width, height)
    return cv2.resize(image, dim, interpolation=cv2.INTER_AREA)


def approximate_contour(contour):
    peri = cv2.arcLength(contour, True)
    return cv2.approxPolyDP(contour, 0.032 * peri, True)


def get_receipt_contour(contours):
    # loop over the contours
    for c in contours:
        approx = approximate_contour(c)
        # if our approximated contour has four points, we can assume it is receipt's rectangle
        if len(approx) == 4:
            return approx


def contour_to_rect(contour, resize_ratio):
    pts = contour.reshape(4, 2)
    rect = np.zeros((4, 2), dtype="float32")
    # top-left point has the smallest sum
    # bottom-right has the largest sum
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]
    rect[2] = pts[np.argmax(s)]
    # compute the difference between the points:
    # the top-right will have the minumum difference
    # the bottom-left will have the maximum difference
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]
    rect[3] = pts[np.argmax(diff)]
    return rect / resize_ratio


def wrap_perspective(img, rect):
    # unpack rectangle points: top left, top right, bottom right, bottom left
    (tl, tr, br, bl) = rect
    # compute the width of the new image
    widthA = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
    widthB = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
    # compute the height of the new image
    heightA = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
    heightB = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
    # take the maximum of the width and height values to reach
    # our final dimensions
    maxWidth = max(int(widthA), int(widthB))
    maxHeight = max(int(heightA), int(heightB))
    # destination points which will be used to map the screen to a "scanned" view
    dst = np.array([
        [0, 0],
        [maxWidth - 1, 0],
        [maxWidth - 1, maxHeight - 1],
        [0, maxHeight - 1]], dtype="float32")
    # calculate the perspective transform matrix
    M = cv2.getPerspectiveTransform(rect, dst)
    # warp the perspective to grab the screen
    return cv2.warpPerspective(img, M, (maxWidth, maxHeight))


# Güncel Receipt Kodu
def getSingleReceiptForTranslate(image_url, lang, is_pdf_or_word):
    if image_url.startswith("data:image/jpeg;base64"):
        url_parts = image_url.split(',')
        base64_image_data = url_parts[1]

        image_data = base64.b64decode(base64_image_data)
        image = Image.open(BytesIO(image_data))

        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    else:
        image_path = image_url.replace("http://127.0.0.1:8000/", "")
        image_cv = cv2.imread(image_path)

    return translatedReceiptOcr(image_cv, lang, is_pdf_or_word)


def translatedReceiptOcr(original_image, lang, is_pdf_or_word):
    extracted_text = pytesseract.image_to_string(original_image)
    if lang != "none":
        translated_text = translate(lang, extracted_text)
        return createDocument(is_pdf_or_word, translated_text.text)
    else:
        return createDocument(is_pdf_or_word, extracted_text)


def getSingleReceiptForSegmentation(image_url, method_type):
    if image_url.startswith("data:image/jpeg;base64"):
        url_parts = image_url.split(',')
        base64_image_data = url_parts[1]

        image_data = base64.b64decode(base64_image_data)
        image = Image.open(BytesIO(image_data))

        image_path = "image.png"
        image.save(image_path)

        image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    else:
        image_path = image_url.replace("http://127.0.0.1:8000/", "")
        image_cv = cv2.imread(image_path)

    if method_type == "default":
        return ocr_for_receipt(image_cv)
    else:
        return receipt_asprise(image_path)


def receipt_asprise(image_path):
    receiptOcrEndpoint = 'https://ocr.asprise.com/api/v1/receipt'
    r = requests.post(receiptOcrEndpoint, data={ \
        'api_key': 'TEST',
        'recognizer': 'auto',
        'ref_no': 'ocr_python_123',
    }, \
                      files={"file": open(image_path, "rb")})

    return createDocument("word", r.text)


def ocr_for_receipt(original_image):
    extracted_text = pytesseract.image_to_string(original_image, config='--psm 6')
    found_lines = find_lines(extracted_text)
    translated_lines = translate_lines("tr", found_lines)
    return create_word_document(translated_lines)


def temp(paths):
    for path in paths:
        original_image = cv2.imread(path)
        extracted_text = pytesseract.image_to_string(original_image, config='--psm 6')
        found_lines = find_lines(extracted_text)

        with open(txtPath, 'a') as file:
            for line in found_lines:
                file.write(line)
                file.write('\n')
            file.write("-------")
            file.write('\n')


def find_lines(metin):
    prices = []
    for line in metin.split('\n'):
        price = re.findall(r'\d+(?:\.\d+|\,\d+)', line)
        if price:
            prices.append(line.strip())
    return prices


def translate(destination, text):
    if text is not None and len(text) > 0:
        translator = Translator()

        translated_text = translator.translate(text, dest=destination)
        return translated_text
    else:
        return ""


def translate_lines(destination, lines):
    if lines is not None and len(lines) > 0:
        translated_lines = []
        for line in lines:
            translator = Translator()
            translated_lines.append(translator.translate(line, dest=destination))
        return translated_lines
    else:
        return ""


def createDocument(docType, text):
    if docType == "text":
        return text
    elif docType == "pdf":
        return create_pdf(text)
    elif docType == "word":
        return create_word(text)


def create_pdf(text):
    pdfkit_config = pdfkit.configuration(wkhtmltopdf='D:/PythonDosyalarim/wkhtml/wkhtmltopdf/bin/wkhtmltopdf.exe')
    html_content = f"<html><body><pre>{text}</pre></body></html>"

    pdf_path = "output.pdf"
    pdfkit.from_string(html_content, "output.pdf", configuration=pdfkit_config)
    return pdf_path


def create_word(text):
    document = Document()
    document.add_paragraph(text)
    docx_path = "output.docx"
    document.save(docx_path)

    return docx_path


def create_word_document(items):
    document = Document()

    document.add_heading('Bulunan Itemlar', level=1)

    for item in items:
        if item is not None:
            if "subtotal" in item.text:
                item.text.replace("subtotal", "Ara Toplam")
            elif "Subtotal" in item.text:
                item.text.replace("Subtotal", "Ara Toplam")
            document.add_paragraph(item.text)

    docx_path = "output.docx"
    document.save(docx_path)

    return docx_path


def parse_receipt(data):
    address, phone = find_address_and_phone_number(data)
    segments = {
        "ID #": None,
        "Walmart": None,
        "(": None,
        "PHONE": phone,
        "ADDRESS": address,
        "Mor": None,
        "Mgr": None,
        "MANAGER": None,
        "ST#": None,
        "ITEMS": parse_items(data),
        "DISCOUNT": None,
        "SUBTOTAL": None,
        "TAX 1": None,
        "TAX 2": None,
        "TAX 3": None,
        "TAX 4": None,
        "TAX 5": None,
        "TOTAL": None,
        "DEBIT TEND": None,
        "CASH TEND": None,
        "CHANGE DUE": None,
        "# ITEMS SOLD": None,
        "TC#": None,
        "EFT DEBIT PAY FROM PRIMARY": None,
        "US DEBIT": None,
        "REF #": None,
        "NETWORK ID.": None,
        "AID": None,
        "AAC": None,
        "TERMINAL": None,
        "01/": None, "02/": None, "03/": None, "04/": None, "05/": None, "06/": None,
        "07/": None, "08/": None, "09/": None, "10/": None, "11/": None, "12/": None,

    }

    for segment, _ in segments.items():
        start_index = data.find(segment)
        if start_index != -1:
            end_index = data.find("\n", start_index)
            segments[segment] = data[start_index:end_index].strip()

    return segments


def parse_items(text):
    items = []
    isItem = False
    isKey = False
    for line in text.split("\n"):
        if (line.startswith("ST#") or line.startswith("S1#") or line.startswith("St#") or line.startswith("SI#")
                or line.startswith("Sl#") or line.startswith("Sı#") or line.startswith("Si#")):
            isKey = True
        elif line.startswith("SUBTOTAL") or line.startswith("DISCOUNT"):
            break
        if isKey and isItem and line != "":
            items.append(line)
        if isKey and not isItem:
            isItem = True

    return items


def find_address_and_phone_number(text):
    address = ""
    phone_number = ""

    lines = text.split('\n')

    walmart_index = text.find("Walmart")
    if walmart_index != -1:
        walmart_lines = text[walmart_index:].split('\n')
        for line in walmart_lines[1:]:
            if re.search(r'\b\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}\b', line):
                phone_number = re.search(r'\b\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}\b', line).group()
                break

    phone_index = text.find(phone_number)
    if phone_index != -1:
        for line in lines:
            if line.startswith(phone_number):
                address_lines = lines[lines.index(line) + 1:]
                for addr_line in address_lines:
                    if addr_line.startswith("ST#"):
                        break
                    address += addr_line.strip() + " "

    return address.strip(), phone_number


"""with open(txtPath, 'w') as f:
    f.close()"""

# findReceipt(r"C:\Users\90553\Desktop\Images\images\9.jpg")
# ocr_for_receipt(cv2.imread(r"C:\Users\90553\Desktop\Images\large-receipt-image-dataset-SRD\1079-receipt.jpg"))
# temp(pathBul(r"C:\Users\90553\Desktop\Images\large-receipt-image-dataset-SRD"))
